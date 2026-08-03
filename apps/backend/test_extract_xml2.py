"""Save full XML of selected paragraphs"""
from docx import Document
from lxml import etree

doc = Document('./uploads/papers/353f79c88c204843b81caf1810895e80.docx')

for i in [12, 13, 14, 22, 23, 24]:
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        xml_str = etree.tostring(para._element, pretty_print=True, encoding='unicode')
        with open(f'./para_{i}_xml.txt', 'w', encoding='utf-8') as f:
            f.write(xml_str)
        print(f"Saved para {i} xml")
