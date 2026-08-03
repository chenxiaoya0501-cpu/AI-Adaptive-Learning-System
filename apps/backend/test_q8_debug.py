"""Debug Q8 extraction step by step"""
from docx import Document
from app.services.word_parser import _extract_images, _extract_para_rich, _extract_options_from_lines
import os, shutil

doc = Document('./uploads/papers/353f79c88c204843b81caf1810895e80.docx')
img_dir = './test_q8_imgs'
if os.path.exists(img_dir):
    shutil.rmtree(img_dir)

image_map = _extract_images(doc, img_dir)

# Q8 likely spans paras 34-37
for i in range(32, 40):
    if i < len(doc.paragraphs):
        rich = _extract_para_rich(doc.paragraphs[i], image_map)
        print(f"Para {i}: {rich[:120]}")

# Q8 lines from current_q_lines would be para 34,35,36,37
q8_lines = [_extract_para_rich(doc.paragraphs[i], image_map) for i in range(34, 38)]
print("\nQ8 lines:")
for i, line in enumerate(q8_lines):
    print(f"  {i}: {line}")

remaining, options = _extract_options_from_lines(q8_lines)
print(f"\nOptions: {options}")
print(f"Remaining: {remaining}")
