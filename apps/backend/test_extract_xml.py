"""Print full XML of paragraphs with missing formulas"""
from docx import Document
from lxml import etree

doc = Document('./uploads/papers/353f79c88c204843b81caf1810895e80.docx')

for i in [12, 13, 14, 22, 23, 24, 40, 41, 42]:
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        print(f"\n{'='*80}\nPara {i}: {para.text[:80]}\n{'='*80}")
        xml_str = etree.tostring(para._element, pretty_print=True, encoding='unicode')
        print(xml_str[:4000])
