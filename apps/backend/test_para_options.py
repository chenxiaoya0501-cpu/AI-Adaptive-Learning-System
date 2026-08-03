"""Debug options paragraphs for Q#8 and Q#3"""
from docx import Document
from lxml import etree

doc = Document('./uploads/papers/353f79c88c204843b81caf1810895e80.docx')

# Q3 options are para 18, 19
# Q8 options are likely around para 30-35
for i in range(30, 42):
    if i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text:
            print(f"\n=== Para {i}: {text[:100]} ===")
            xml_str = etree.tostring(para._element, encoding='unicode')
            print(xml_str[:800])
