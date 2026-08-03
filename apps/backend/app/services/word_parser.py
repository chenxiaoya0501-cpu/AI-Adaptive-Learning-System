"""试卷Word文档解析服务 - 从Word文件中提取结构化题目数据

核心思路：
1. 遍历docx段落XML，提取每段的富文本（文字 + [IMG:xxx] 图片占位符）
2. 过滤非题目内容（试卷标题、答题须知、注意事项等）
3. 按题号分割题目，识别大题类型（选择/填空/解答）
4. 从大题说明读取「每小题 X 分」或「17-21每题8分，24题12分」题号分值表；
   题干「（X分）」优先，写入每题 score
5. 将【分析】【解答】【点评】等标记的内容分离为 answer / analysis 字段
6. 提取选择题选项（支持选项含图片）
"""
import os
import re
import logging
from typing import List, Dict, Any, Optional, Tuple
from lxml import etree

from docx import Document

logger = logging.getLogger(__name__)

# XML namespaces
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
V_NS = 'urn:schemas-microsoft-com:vml'  # VML (old formula images)

# 非题目内容的关键词（用于过滤试卷头部信息）
SKIP_KEYWORDS = [
    '答题前', '务必', '姓名', '准考证', '考生须知', '注意事项',
    '本试题卷', '本试卷', '考试时间', '满分', '答题卡',
    '密封线', '装订线', '机密', '启用前', '试卷类型',
    '选择题部分', '非选择题部分',
    '参考答案与试题解析',
    '答题时', '黑色字迹', '钢笔', '签字笔', '2B铅笔',
    '草稿纸', '连接', '连结', '同义',
]

# 大题类型映射（按子串匹配；「单选题」不含「选择」二字，需单独映射）
SECTION_TYPE_MAP = {
    '单选': 'choice',
    '多选': 'choice',
    '选择': 'choice',
    '填空': 'fill',
    '解答': 'answer',
    '计算': 'answer',
    '证明': 'proof',
    '问答': 'answer',
    '应用': 'answer',
    '综合': 'answer',
}

# 题号正则（匹配行首的题号）
QUESTION_NUM_RE = re.compile(
    r'^(\d{1,2})\s*[．.、]\s*'  # 1．/ 1. / 1、
    r'|^第\s*(\d{1,2})\s*题\s*'  # 第1题
)

# 大题标记正则
SECTION_RE = re.compile(
    r'^[一二三四五六七八九十]+\s*[、．.]\s*(.+?)[\s（(（]'
)

# 选项正则（支持行内多选项，如 A．xxx  B．xxx，选项值可以是图片）
# 也支持 A xxx / B[IMG...]（选项字母后直接跟图片占位符，无分隔符）
OPTION_LINE_RE = re.compile(
    r'([A-D])(?:\s*[．.、)）]\s*|[ \t]+(?=\S)|(?=\[))'
)

# 答案标记（简短答案，如"A"、"故选C"）
ANSWER_MARKERS = ['【答案】', '【答】']
# 解析标记（详细解题过程）
ANALYSIS_MARKERS = ['【解析】', '【解答】', '【分析】', '【详解】', '【思路】']
# 评注标记（杭州卷常用「点睛」）
COMMENT_MARKERS = ['【点评】', '【评注】', '【点睛】']

# 上下分卷：正文题目结束后的答案附录起始标记（整行或行首）
APPENDIX_START_MARKERS = (
    '参考答案与试题解析',
    '参考答案与解析',
    '答案与试题解析',
    '答案与解析',
    '试题解析',
    '参考答案',
)


def _convert_doc_to_docx(doc_path: str) -> str:
    """将.doc格式转换为.docx格式（仅Windows，使用Word COM）"""
    docx_path = doc_path + "x"
    if os.path.exists(docx_path):
        return docx_path
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
        doc.Close()
        word.Quit()
        logger.info(f"已将 .doc 转换为 .docx: {docx_path}")
        return docx_path
    except ImportError:
        raise RuntimeError("无法解析.doc格式文件：需要安装pywin32或将文件另存为.docx格式后重新上传")
    except Exception as e:
        raise RuntimeError(f"将.doc转换为.docx失败: {e}。请将文件另存为.docx格式后重新上传")


def parse_exam_word(file_path: str, image_save_dir: str) -> List[Dict[str, Any]]:
    """解析试卷Word文档，提取结构化题目列表

    Args:
        file_path: Word文件路径（支持.doc和.docx）
        image_save_dir: 图片保存目录

    Returns:
        题目列表，每个题目为字典，包含 question_number, question_type,
        content, options, answer, analysis, difficulty, images
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    if file_path.lower().endswith('.doc') and not file_path.lower().endswith('.docx'):
        file_path = _convert_doc_to_docx(file_path)

    os.makedirs(image_save_dir, exist_ok=True)
    doc = Document(file_path)

    # 0. 先扫描所有段落XML，收集每个rId对应的VML显示尺寸(pt)
    vml_sizes = _collect_vml_sizes(doc)

    # 1. 提取所有图片到磁盘，建立 rId -> filename 映射
    image_map = _extract_images(doc, image_save_dir, vml_sizes)

    # 2. 提取所有段落的富文本（含图片占位符）
    rich_lines = _extract_rich_paragraphs(doc, image_map)

    # 3. 过滤、分割、结构化
    questions = _parse_questions_from_lines(rich_lines)

    logger.info(f"Word解析完成：共提取 {len(questions)} 道题目")
    return questions


# ==================== 图片提取 ====================

def _collect_vml_sizes(doc: Document) -> Dict[str, Tuple[float, float]]:
    """扫描文档所有段落，收集VML v:shape中每个图片rId对应的显示尺寸(pt)
    
    返回 {rId: (width_pt, height_pt)}
    """
    sizes = {}
    for para in doc.paragraphs:
        for elem in para._element.iter():
            tag = etree.QName(elem.tag).localname if '}' in elem.tag else elem.tag
            ns = etree.QName(elem.tag).namespace if '}' in elem.tag else ''
            if tag == 'imagedata' and ns == V_NS:
                rid = elem.get(f'{{{R_NS}}}id')
                if rid and rid not in sizes:
                    parent = elem.getparent()
                    if parent is not None:
                        style_str = parent.get('style', '')
                        w_pt, h_pt = _parse_vml_style_size(style_str)
                        if w_pt > 0 and h_pt > 0:
                            sizes[rid] = (w_pt, h_pt)
    return sizes


def _extract_images(doc: Document, save_dir: str, vml_sizes: Dict[str, Tuple[float, float]] = None) -> Dict[str, str]:
    """仏Word文档中提取所有图片，返回 {rId: 图片文件名}
    
    WMF/EMF格式会自动转换为PNG（浏览器不支持WMF/EMF）
    vml_sizes: 从VML中收集的 {rId: (width_pt, height_pt)}，用于正确转换宽高比
    """
    if vml_sizes is None:
        vml_sizes = {}
    os.makedirs(save_dir, exist_ok=True)
    image_map = {}
    counter = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            counter += 1
            ct = rel.target_part.content_type
            blob = rel.target_part.blob
            needs_convert = False

            if "jpeg" in ct or "jpg" in ct:
                ext = ".jpg"
            elif "gif" in ct:
                ext = ".gif"
            elif "png" in ct:
                ext = ".png"
            elif "wmf" in ct or "emf" in ct:
                ext = ".png"
                needs_convert = True
            elif "bmp" in ct:
                ext = ".png"
                needs_convert = True
            else:
                ext = ".png"

            filename = f"img_{counter:03d}{ext}"
            filepath = os.path.join(save_dir, filename)

            if needs_convert:
                # 使用VML中的原始pt尺寸计算目标像素尺寸
                target_w, target_h = None, None
                if rel.rId in vml_sizes:
                    w_pt, h_pt = vml_sizes[rel.rId]
                    # 1pt ≈ 1.333px，但我们用 2x 以保证清晰度
                    target_w = int(w_pt * 2)
                    target_h = int(h_pt * 2)
                converted = _convert_image_to_png(blob, filepath, target_w, target_h)
                if not converted:
                    with open(filepath, "wb") as f:
                        f.write(blob)
            else:
                # PNG/JPG/GIF：如果有VML尺寸信息，按目标2x尺寸缩放以保证清晰
                if rel.rId in vml_sizes:
                    w_pt, h_pt = vml_sizes[rel.rId]
                    target_w = int(w_pt * 2)
                    target_h = int(h_pt * 2)
                    _resize_image_to_target(blob, filepath, target_w, target_h)
                else:
                    with open(filepath, "wb") as f:
                        f.write(blob)

            image_map[rel.rId] = filename
    logger.info(f"提取到 {counter} 张图片")
    return image_map


def _resize_image_to_target(blob: bytes, output_path: str, target_w: int, target_h: int):
    """将PNG/JPG图片按目标尺寸缩放保存，保证清晰度
    
    如果原始图片分辨率已经 >= 目标尺寸，直接保存原始数据。
    否则按目标比例缩放。
    """
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(blob))
        orig_w, orig_h = img.size
        # 只有在原始尺寸小于目标时才放大（避免无谓缩小损失质量）
        if orig_w < target_w or orig_h < target_h:
            img = img.resize((target_w, target_h), Image.LANCZOS)
            img.save(output_path, 'PNG')
        else:
            # 原始分辨率足够，直接保存
            with open(output_path, "wb") as f:
                f.write(blob)
    except Exception as e:
        logger.warning(f"图片缩放失败，直接保存原始: {e}")
        with open(output_path, "wb") as f:
            f.write(blob)


def _convert_image_to_png(blob: bytes, output_path: str, target_w: int = None, target_h: int = None) -> bool:
    """将WMF/EMF/BMP图片数据转换为PNG格式
    
    优先使用Windows GDI+进行高质量矢量渲染（适用于公式WMF）。
    如果GDI+不可用，回退到Pillow。
    
    Args:
        blob: 原始图片二进制数据
        output_path: 输出PNG路径
        target_w: 目标宽度(像素)，基于VML style中的pt值换算
        target_h: 目标高度(像素)，基于VML style中的pt值换算
    
    Returns:
        True if conversion succeeded, False otherwise
    """
    import tempfile
    # 先保存到临时WMF文件
    with tempfile.NamedTemporaryFile(suffix='.wmf', delete=False) as tmp:
        tmp.write(blob)
        tmp_path = tmp.name
    
    try:
        # 优先使用GDI+高质量渲染
        if _convert_wmf_gdiplus(tmp_path, output_path, target_w, target_h):
            return True
        # 回退到Pillow
        return _convert_wmf_pillow(tmp_path, output_path, target_w, target_h)
    finally:
        os.unlink(tmp_path)


def _convert_wmf_gdiplus(wmf_path: str, output_path: str, target_w: int = None, target_h: int = None) -> bool:
    """使用Windows GDI+将WMF/EMF渲染为高质量PNG（矢量渲染，公式清晰）"""
    try:
        import ctypes
        gdiplus = ctypes.windll.gdiplus
        
        class GdiplusStartupInput(ctypes.Structure):
            _fields_ = [
                ("GdiplusVersion", ctypes.c_uint32),
                ("DebugEventCallback", ctypes.c_void_p),
                ("SuppressBackgroundThread", ctypes.c_int),
                ("SuppressExternalCodecs", ctypes.c_int),
            ]
        
        class RectF(ctypes.Structure):
            _fields_ = [("X", ctypes.c_float), ("Y", ctypes.c_float),
                       ("Width", ctypes.c_float), ("Height", ctypes.c_float)]
        
        class GUID(ctypes.Structure):
            _fields_ = [("Data1", ctypes.c_ulong), ("Data2", ctypes.c_ushort),
                       ("Data3", ctypes.c_ushort), ("Data4", ctypes.c_ubyte * 8)]
        
        # 初始化GDI+
        startup_input = GdiplusStartupInput()
        startup_input.GdiplusVersion = 1
        token = ctypes.c_ulong()
        gdiplus.GdiplusStartup(ctypes.byref(token), ctypes.byref(startup_input), None)
        
        try:
            # 加载WMF/EMF元文件
            metafile = ctypes.c_void_p()
            status = gdiplus.GdipCreateMetafileFromFile(
                ctypes.c_wchar_p(os.path.abspath(wmf_path)),
                ctypes.byref(metafile)
            )
            if status != 0:
                return False
            
            # 获取元文件原始尺寸
            width = ctypes.c_float()
            height = ctypes.c_float()
            gdiplus.GdipGetImageDimension(metafile, ctypes.byref(width), ctypes.byref(height))
            
            # 计算目标渲染尺寸
            if target_w and target_h and target_w > 0 and target_h > 0:
                # 使用VML中指定的尺寸，再放大2倍保证清晰
                w = target_w * 2
                h = target_h * 2
            else:
                # 使用元文件原始尺寸的合适缩放
                # GDI+报告的尺寸单位是0.01mm，需转换为合理像素
                # 使用scale使得结果在合适范围内（100-600px）
                meta_w = width.value
                meta_h = height.value
                # 目标：短边至少60px，长边不超过600px
                scale = max(60 / min(meta_w, meta_h), 1.0)
                scale = min(scale, 600 / max(meta_w, meta_h))
                w = max(int(meta_w * scale), 30)
                h = max(int(meta_h * scale), 30)
            
            # 创建位图
            bitmap = ctypes.c_void_p()
            # PixelFormat32bppARGB = 0x26200A
            gdiplus.GdipCreateBitmapFromScan0(w, h, 0, 0x26200A, None, ctypes.byref(bitmap))
            
            # 从位图创建Graphics对象
            graphics = ctypes.c_void_p()
            gdiplus.GdipGetImageGraphicsContext(bitmap, ctypes.byref(graphics))
            
            # 设置高质量渲染
            gdiplus.GdipSetSmoothingMode(graphics, 2)  # HighQuality
            gdiplus.GdipSetTextRenderingHint(graphics, 1)  # AntiAlias
            gdiplus.GdipSetInterpolationMode(graphics, 7)  # HighQualityBicubic
            
            # 白色背景
            gdiplus.GdipGraphicsClear(graphics, 0xFFFFFFFF)
            
            # 绘制元文件
            gdiplus.GdipDrawImageRectI(graphics, metafile, 0, 0, w, h)
            
            # 保存为PNG
            png_clsid = GUID(0x557cf406, 0x1a04, 0x11d3,
                             (ctypes.c_ubyte * 8)(0x9a, 0x73, 0x00, 0x00, 0xf8, 0x1e, 0xf3, 0x2e))
            
            status = gdiplus.GdipSaveImageToFile(
                bitmap, ctypes.c_wchar_p(os.path.abspath(output_path)),
                ctypes.byref(png_clsid), None
            )
            
            # 清理
            gdiplus.GdipDisposeImage(metafile)
            gdiplus.GdipDisposeImage(bitmap)
            gdiplus.GdipDeleteGraphics(graphics)
            
            return status == 0
        finally:
            gdiplus.GdiplusShutdown(token)
    except Exception as e:
        logger.debug(f"GDI+渲染WMF失败，将回退到Pillow: {e}")
        return False


def _convert_wmf_pillow(wmf_path: str, output_path: str, target_w: int = None, target_h: int = None) -> bool:
    """使用Pillow将WMF转PNG（回退方案，质量较低）"""
    try:
        from PIL import Image
        img = Image.open(wmf_path)
        if target_w and target_h and target_w > 0 and target_h > 0:
            img = img.resize((target_w * 2, target_h * 2), Image.LANCZOS)
        else:
            w, h = img.size
            if w < 100 or h < 100:
                scale = max(100 / w, 100 / h, 2)
                img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        img.save(output_path, 'PNG')
        return True
    except Exception as e:
        logger.warning(f"Pillow转换PNG失败: {e}")
        return False


def _parse_vml_style_size(style_str: str) -> Tuple[float, float]:
    """从VML shape的style属性中解析宽高(pt)
    
    例如: 'height:30.75pt;width:12pt;' -> (12.0, 30.75)
    返回 (width_pt, height_pt)，解析失败返回 (0, 0)
    """
    w_pt, h_pt = 0.0, 0.0
    if not style_str:
        return (0, 0)
    # 匹配 width:Xpt 和 height:Xpt
    w_match = re.search(r'width:\s*([\d.]+)\s*pt', style_str)
    h_match = re.search(r'height:\s*([\d.]+)\s*pt', style_str)
    if w_match:
        w_pt = float(w_match.group(1))
    if h_match:
        h_pt = float(h_match.group(1))
    return (w_pt, h_pt)


# ==================== 富文本段落提取 ====================

WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'


def _get_drawing_size_pt(blip_elem) -> Tuple[float, float]:
    """从DrawingML a:blip元素向上查找wp:extent获取图片尺寸(pt)
    
    EMU (English Metric Units): 1pt = 12700 EMU
    返回 (width_pt, height_pt)
    """
    EMU_PER_PT = 12700
    # 向上遍历祖先节点查找wp:inline或wp:anchor
    parent = blip_elem.getparent()
    while parent is not None:
        ptag = etree.QName(parent.tag).localname if '}' in parent.tag else parent.tag
        pns = etree.QName(parent.tag).namespace if '}' in parent.tag else ''
        if ptag in ('inline', 'anchor') and pns == WP_NS:
            # 查找wp:extent子元素
            for child in parent:
                ctag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
                if ctag == 'extent':
                    cx = child.get('cx', '0')
                    cy = child.get('cy', '0')
                    try:
                        w_pt = int(cx) / EMU_PER_PT
                        h_pt = int(cy) / EMU_PER_PT
                        return (w_pt, h_pt)
                    except (ValueError, TypeError):
                        pass
            break
        parent = parent.getparent() if hasattr(parent, 'getparent') else None
    return (0.0, 0.0)


def _extract_para_rich(para, image_map: Dict[str, str]) -> str:
    """从单个段落的XML中提取富文本，图片用 [IMG:filename] 占位
    
    支持：
    - w:t 文字
    - w:tab 制表符
    - w:br 换行
    - a:blip 图片（DrawingML）
    - v:imagedata 图片（VML，老版本公式图片）
    - m:oMath / m:oMathPara 数学公式文字
    """
    parts = []
    added_rids = set()  # 防止同一个图片被重复添加

    for elem in para._element.iter():
        tag = etree.QName(elem.tag).localname if '}' in elem.tag else elem.tag
        ns = etree.QName(elem.tag).namespace if '}' in elem.tag else ''

        # 文本
        if tag == 't' and ns == W_NS:
            # 如果 w:t 位于 m:oMath 内部，等下统一用 oMath 处理；避免重复
            parent = elem
            in_math = False
            while parent is not None:
                if etree.QName(parent.tag).namespace == M_NS:
                    in_math = True
                    break
                parent = parent.getparent() if hasattr(parent, 'getparent') else None
            if not in_math and elem.text:
                parts.append(elem.text)

        elif tag == 'tab' and ns == W_NS:
            parts.append('    ')

        elif tag == 'br' and ns == W_NS:
            parts.append('\n')

        # 数学公式文字（直接作为文字保留）
        elif tag in ('oMath', 'oMathPara') and ns == M_NS:
            math_text = ''.join(elem.itertext())
            if math_text.strip():
                parts.append(math_text)

        # VML图片（老公式嵌入形式：<v:imagedata r:id="..."/>）
        elif tag == 'imagedata' and ns == V_NS:
            # VML使用 r:id 而不是 r:embed
            rid = elem.get(f'{{{R_NS}}}id')
            if rid and rid not in added_rids and rid in image_map:
                # 从父级v:shape的style中读取原始宽高
                w_pt, h_pt = 0.0, 0.0
                parent = elem.getparent()
                if parent is not None:
                    style_str = parent.get('style', '')
                    w_pt, h_pt = _parse_vml_style_size(style_str)
                if w_pt > 0 and h_pt > 0:
                    parts.append(f'[IMG:{image_map[rid]},{w_pt:.1f},{h_pt:.1f}]')
                else:
                    parts.append(f'[IMG:{image_map[rid]}]')
                added_rids.add(rid)

        # DrawingML图片（<a:blip r:embed="..."/>）
        elif tag == 'blip' and ns == A_NS:
            rid = elem.get(f'{{{R_NS}}}embed')
            if rid and rid not in added_rids and rid in image_map:
                # 尝试从祖先wp:inline/wp:anchor的wp:extent读取尺寸(EMU)
                w_pt, h_pt = _get_drawing_size_pt(elem)
                if w_pt > 0 and h_pt > 0:
                    parts.append(f'[IMG:{image_map[rid]},{w_pt:.1f},{h_pt:.1f}]')
                else:
                    parts.append(f'[IMG:{image_map[rid]}]')
                added_rids.add(rid)

    return ''.join(parts)


def _extract_rich_paragraphs(doc: Document, image_map: Dict[str, str]) -> List[str]:
    """提取文档所有段落的富文本列表"""
    lines = []
    for para in doc.paragraphs:
        rich = _extract_para_rich(para, image_map)
        lines.append(rich)
    return lines


# ==================== 题目解析 ====================

def _is_skip_line(line: str) -> bool:
    """判断是否为应跳过的非题目内容（答题须知等）"""
    text = line.strip()
    if not text:
        return True
    for kw in SKIP_KEYWORDS:
        if kw in text:
            return True
    # 纯数字行（如页码）
    if re.match(r'^\d{1,3}$', text):
        return True
    return False


def _extract_score_per_question(text: str) -> Optional[float]:
    """从大题说明中提取统一的「每小题/每题 X 分」。

    例：一、选择题（本大题共10小题，每小题3分，共30分）→ 3.0
    仅有「共30分」而无每题分时返回 None（不强行均分）。
    若说明中已是「17-21每题8分」这类题号表，应走 _extract_per_number_scores，勿用本函数首个「每题」。
    """
    if not text:
        return None
    # 含题号区间/列举分值表时，不作统一每题分
    if _extract_per_number_scores(text):
        return None
    patterns = [
        r'每小题\s*(\d+(?:\.\d+)?)\s*分',
        r'每题\s*(\d+(?:\.\d+)?)\s*分',
        r'小题各\s*(\d+(?:\.\d+)?)\s*分',
        r'各\s*(\d+(?:\.\d+)?)\s*分',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            try:
                return float(m.group(1))
            except (TypeError, ValueError):
                return None
    return None


def _extract_per_number_scores(text: str) -> Dict[int, float]:
    """从大题说明解析按题号分值表。

    例：三、解答题（17-21 每题 8 分，22、23 每题 10 分，24 题 12 分）
    → {17:8,...,21:8, 22:10, 23:10, 24:12}
    """
    result: Dict[int, float] = {}
    if not text:
        return result

    # 区间：17-21 每题 8 分
    for m in re.finditer(
        r'(\d{1,2})\s*[-~～—至到]\s*(\d{1,2})\s*每题\s*(\d+(?:\.\d+)?)\s*分',
        text,
    ):
        a, b, s = int(m.group(1)), int(m.group(2)), float(m.group(3))
        for n in range(min(a, b), max(a, b) + 1):
            result[n] = s

    # 列举：22、23 每题 10 分
    for m in re.finditer(
        r'((?:\d{1,2}\s*[、，,]\s*)+\d{1,2})\s*每题\s*(\d+(?:\.\d+)?)\s*分',
        text,
    ):
        nums = [int(x) for x in re.findall(r'\d{1,2}', m.group(1))]
        s = float(m.group(2))
        for n in nums:
            result[n] = s

    # 单题：24 题 12 分 / 第24题12分（避免误伤「小题」）
    for m in re.finditer(
        r'(?:第\s*)?(\d{1,2})\s*题\s*(\d+(?:\.\d+)?)\s*分',
        text,
    ):
        # 「10小题」不会匹配；「每小题3分」前无独立题号
        idx = m.start(1)
        if idx > 0 and text[idx - 1] in ('小', '大'):
            continue
        result[int(m.group(1))] = float(m.group(2))

    return result


def _extract_score_from_stem(text: str) -> Optional[float]:
    """题干开头「（8分）」/「(8分)」。"""
    if not text:
        return None
    m = re.match(
        r'^[（(]\s*(\d+(?:\.\d+)?)\s*分\s*[）)]',
        text.strip(),
    )
    if not m:
        return None
    try:
        return float(m.group(1))
    except (TypeError, ValueError):
        return None


def _detect_section(line: str) -> Tuple[Optional[str], Optional[float], Dict[int, float]]:
    """检测大题标记，返回 (题目类型, 统一每题分, 题号分值表)。"""
    text = line.strip()
    m = SECTION_RE.match(text)
    if not m:
        # 也匹配不带括号的，如 "一、选择题" 或整行含说明
        m2 = re.match(r'^[一二三四五六七八九十]+\s*[、．.]\s*(.+)', text)
        if m2:
            section_name = m2.group(1)
        else:
            return None, None, {}
    else:
        section_name = m.group(1)

    qtype = None
    for key, mapped in SECTION_TYPE_MAP.items():
        if key in section_name or key in text:
            qtype = mapped
            break
    if qtype is None:
        return None, None, {}
    per_number = _extract_per_number_scores(text)
    score_each = None if per_number else _extract_score_per_question(text)
    return qtype, score_each, per_number


def _detect_question_num(line: str) -> Optional[int]:
    """检测题号，返回题号数字或None"""
    m = QUESTION_NUM_RE.match(line.strip())
    if m:
        num_str = m.group(1) or m.group(2)
        return int(num_str)
    return None


def _strip_question_num(line: str) -> str:
    """去掉行首的题号部分"""
    m = QUESTION_NUM_RE.match(line.strip())
    if m:
        return line.strip()[m.end():].strip()
    return line.strip()


def _extract_options_from_lines(lines: List[str]) -> Tuple[List[str], Dict[str, str]]:
    """从内容行中提取选项，返回 (剩余行, 选项字典)

    支持两种格式：
    - 每行一个选项：A．xxx 或 A．[IMG:xxx]
    - 一行多个选项：A．xxx  B．xxx  C．xxx  D．xxx
    - 选项值可以是纯图片
    """
    options = {}
    remaining = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 检查是否为选项行（以A/B/C/D开头）
        # 先尝试拆分一行中的多个选项
        opt_positions = list(OPTION_LINE_RE.finditer(stripped))
        if opt_positions and len(opt_positions) >= 2 and opt_positions[0].start() == 0:
            # 一行多个选项（第一个选项必须在行首）
            for idx, m in enumerate(opt_positions):
                letter = m.group(1)
                start = m.end()
                end = opt_positions[idx + 1].start() if idx + 1 < len(opt_positions) else len(stripped)
                val = stripped[start:end].strip()
                options[letter] = val if val else ''
        elif opt_positions and len(opt_positions) == 1:
            m = opt_positions[0]
            # 只有当选项字母在行首（位置0）时才算选项行
            if m.start() == 0:
                letter = m.group(1)
                val = stripped[m.end():].strip()
                options[letter] = val if val else ''
            else:
                remaining.append(line)
        else:
            remaining.append(line)

    return remaining, options


def _split_answer_analysis(content: str) -> Tuple[str, Optional[str], Optional[str]]:
    """从题目内容中分离出答案和解析

    标记优先级：
    - 【答案】 → answer字段（简短答案，如"A"、"x=3"）
    - 【解析】/【解答】/【分析】 → analysis字段（详细解题过程）
    - 【点评】 → 追加到analysis

    返回：(纯题目内容, 答案, 解析)
    """
    all_markers = ANSWER_MARKERS + ANALYSIS_MARKERS + COMMENT_MARKERS

    # 找到第一个标记的位置
    first_marker_pos = len(content)
    for marker in all_markers:
        pos = content.find(marker)
        if pos != -1 and pos < first_marker_pos:
            first_marker_pos = pos

    if first_marker_pos == len(content):
        # 没有任何标记，尝试从内容中提取简短答案（如"故选：C"）
        short = _extract_short_answer(content)
        if short:
            return content.strip(), short, None
        return content.strip(), None, None

    pure_content = content[:first_marker_pos].strip()
    marker_section = content[first_marker_pos:]

    answer = None
    analysis = None

    # 提取【答案】部分 → answer
    for marker in ANSWER_MARKERS:
        pos = marker_section.find(marker)
        if pos != -1:
            end = len(marker_section)
            for m2 in ANALYSIS_MARKERS + COMMENT_MARKERS:
                p2 = marker_section.find(m2, pos + len(marker))
                if p2 != -1 and p2 < end:
                    end = p2
            answer_text = marker_section[pos + len(marker):end].strip()
            if answer_text:
                answer = answer_text
            break

    # 提取【解析】/【解答】/【分析】部分 → analysis
    for marker in ANALYSIS_MARKERS:
        pos = marker_section.find(marker)
        if pos != -1:
            end = len(marker_section)
            for m2 in ANSWER_MARKERS + COMMENT_MARKERS:
                p2 = marker_section.find(m2, pos + len(marker))
                if p2 != -1 and p2 < end:
                    end = p2
            analysis_text = marker_section[pos + len(marker):end].strip()
            if analysis_text:
                analysis = analysis_text
            break

    # 如果没有analysis，把【点评】当作analysis
    if not analysis:
        for marker in COMMENT_MARKERS:
            pos = marker_section.find(marker)
            if pos != -1:
                analysis_text = marker_section[pos + len(marker):].strip()
                if analysis_text:
                    analysis = analysis_text
                break

    # 如果有analysis但没有answer，尝试从analysis中提取简短答案
    if analysis and not answer:
        short = _extract_short_answer(analysis)
        if short:
            answer = short

    # 如果没有单独的【答案】标记但有【解答】内容包含完整解答过程
    # 从中提取简短答案
    if answer and not analysis:
        # answer可能是完整解答过程（来自【解答】），拆分
        if len(answer) > 50:
            short = _extract_short_answer(answer)
            if short:
                analysis = answer
                answer = short

    return pure_content, answer, analysis


def _extract_short_answer(answer_text: str) -> Optional[str]:
    """从解答过程中提取简短答案（如 故选：C、故答案为xxx）

    注意：答案可能是公式图片占位符 [IMG:file.png,W,H]。
    终止符只用中文句号「．」「。」或换行，避免 ASCII「.」截断 .png 文件名。
    """
    m = re.search(r'故选[：:]\s*([A-D])', answer_text)
    if m:
        return m.group(1).strip()

    for prefix in (r'故答案为[：:]?\s*', r'答案[是为][：:]?\s*'):
        m = re.search(prefix, answer_text)
        if not m:
            continue
        rest = answer_text[m.end():]
        # 公式图片：完整吃掉 [IMG:...]（文件名含 .png 等）
        img_m = re.match(r'(\[IMG:[^\]]+\])', rest)
        if img_m:
            return img_m.group(1).strip()
        # 纯文本：不以 ASCII「.」收尾（避免与小数/扩展名冲突）
        text_m = re.match(r'(.+?)(?:[．。\n]|$)', rest, flags=re.DOTALL)
        if text_m:
            val = text_m.group(1).strip()
            if val:
                return val
    return None


def _collect_question_images(content: str, options: Dict[str, str]) -> List[str]:
    """收集题目中引用的所有图片文件名"""
    images = []
    pattern = re.compile(r'\[IMG:([^\],]+)')
    for m in pattern.finditer(content):
        images.append(m.group(1))
    for v in (options or {}).values():
        for m in pattern.finditer(v):
            images.append(m.group(1))
    return images


def _is_appendix_start_line(text: str) -> bool:
    """是否为「上题下答」分卷的答案附录标题行。"""
    t = (text or "").strip()
    if not t:
        return False
    for m in APPENDIX_START_MARKERS:
        if t == m or t.startswith(m + " ") or t.startswith(m + "（") or t.startswith(m + "("):
            return True
    return False


def _find_split_appendix_index(rich_lines: List[str]) -> Optional[int]:
    """定位答案附录起始行。

    杭州 2021 等「解析卷」常见结构：
    1) 上半：一、单选题 / 二、填空 / 三、解答（仅题干）
    2) 中间：独立一行「参考答案」
    3) 下半：按题号重复 1．B +【分析】【详解】【点睛】

    优先认附录标题；否则用题号回绕（如 23 后出现 1．B）兜底。
    注意：考生须知里的「1. 2. 3.」不得参与题号统计，否则会把真正第1题误判为附录起点。
    """
    started = False
    last_q = 0
    max_q = 0
    for i, line in enumerate(rich_lines):
        stripped = (line or "").strip()
        if not stripped:
            continue
        # 须已进入大题区，且已收到若干真题后，才认「参考答案」标题
        if started and max_q >= 3 and _is_appendix_start_line(stripped):
            return i
        sec_type, _, _ = _detect_section(stripped)
        if sec_type is not None:
            started = True
            continue
        q_num = _detect_question_num(stripped)
        if q_num is None:
            continue
        # 大题标记之前的编号（考生须知）一律忽略
        if not started:
            continue
        # 题号明显回绕（已收到至少 3 题后出现更小的题号）→ 答案附录
        if max_q >= 3 and q_num < last_q and q_num <= 5:
            return i
        last_q = q_num
        max_q = max(max_q, q_num)
    return None


def _attach_appendix_line(questions: List[Dict[str, Any]], q_num: int, body: str) -> None:
    """把答案附录中的一行合并到已解析题目上（按题号）。"""
    target = next((q for q in questions if q.get("question_number") == q_num), None)
    if not target:
        return
    body = (body or "").strip()
    if not body:
        return
    letter = re.match(r'^([A-D])\s*[．.。]?$', body)
    if letter and not target.get("answer"):
        target["answer"] = letter.group(1)
        return
    # 解析/点评起笔 → 进 analysis
    if re.match(r'^(解[：:．.]|试题|本题|分析|详解|【)', body):
        prev = target.get("analysis") or ""
        target["analysis"] = f"{prev}\n{body}".strip() if prev else body
        return
    pure, ans, analysis = _split_answer_analysis(body)
    if ans and not target.get("answer"):
        target["answer"] = ans
    chunk = analysis or pure
    if chunk and ans and chunk.strip() in {ans, f"故选：{ans}", f"故选:{ans}"}:
        return
    if chunk and (analysis or any(m in body for m in ANALYSIS_MARKERS + COMMENT_MARKERS)):
        prev = target.get("analysis") or ""
        target["analysis"] = f"{prev}\n{chunk}".strip() if prev else chunk
        return
    # 短正文且尚无答案：填空附录常见「3．1」「12．5a」
    if not target.get("answer") and len(body) <= 80:
        target["answer"] = body.rstrip("．.。").strip()
        return
    if chunk:
        prev = target.get("analysis") or ""
        target["analysis"] = f"{prev}\n{chunk}".strip() if prev else chunk


def _merge_answer_appendix_lines(
    questions: List[Dict[str, Any]],
    appendix_lines: List[str],
) -> None:
    """第二阶段：解析答案附录，按题号写回 answer/analysis（绝不新建题目）。

    先按题号收集所有附录行，合并为完整文本块后，再提取答案/解析。
    这样可以正确处理附录中「先重复题干 → 再接【答案】【解析】」的格式。
    """
    # ---- 第 1 步：按题号收集附录行 ----
    collected: List[Tuple[int, List[str]]] = []  # [(q_num, [lines])]
    current_q: Optional[int] = None
    current_lines: List[str] = []

    def _flush_collected():
        nonlocal current_q, current_lines
        if current_q is not None and current_lines:
            collected.append((current_q, current_lines))
        current_lines = []

    for line in appendix_lines:
        stripped = (line or "").strip()
        if not stripped:
            continue
        if _is_appendix_start_line(stripped):
            continue
        if _detect_section(stripped)[0] is not None:
            continue
        q_num = _detect_question_num(stripped)
        if q_num is not None:
            _flush_collected()
            current_q = q_num
            after = _strip_question_num(stripped)
            if after:
                current_lines.append(after)
            continue
        if current_q is not None and not _is_skip_line(stripped):
            current_lines.append(stripped)
    _flush_collected()

    # ---- 第 2 步：对每个题号的合并文本提取答案/解析 ----
    all_markers = ANSWER_MARKERS + ANALYSIS_MARKERS + COMMENT_MARKERS
    for q_num, lines in collected:
        target = next((q for q in questions if q.get("question_number") == q_num), None)
        if not target:
            continue
        merged = '\n'.join(lines).strip()
        if not merged:
            continue

        # 如果合并文本中包含【答案】或【解析】标记，走标记提取逻辑
        has_marker = any(m in merged for m in all_markers)
        if has_marker:
            _, ans, analysis = _split_answer_analysis(merged)
            if ans and not target.get("answer"):
                target["answer"] = ans
            if analysis:
                prev = target.get("analysis") or ""
                target["analysis"] = f"{prev}\n{analysis}".strip() if prev else analysis
            continue

        # 无标记：逐行走旧的 _attach_appendix_line 逻辑（填空题简答等）
        for body in lines:
            _attach_appendix_line(questions, q_num, body)


def _parse_stem_region(rich_lines: List[str]) -> List[Dict[str, Any]]:
    """第一阶段：只解析题干区（不含答案附录）。"""
    questions: List[Dict[str, Any]] = []
    current_section_type = None
    current_section_score: Optional[float] = None
    current_section_per_number: Dict[int, float] = {}
    current_q_num = None
    current_q_lines: List[str] = []
    started = False
    last_q_num = 0

    def _resolve_question_score(q_num: int, content: str) -> Optional[float]:
        stem = _extract_score_from_stem(content)
        if stem is not None:
            return stem
        if q_num in current_section_per_number:
            return current_section_per_number[q_num]
        return current_section_score

    def _flush_question():
        nonlocal current_q_num, current_q_lines
        if current_q_num is None or not current_q_lines:
            current_q_lines = []
            return

        raw_content = '\n'.join(current_q_lines).strip()
        if not raw_content:
            current_q_lines = []
            return

        content_lines_split = raw_content.split('\n')
        remaining_lines, options = _extract_options_from_lines(content_lines_split)
        content_no_opts = '\n'.join(remaining_lines).strip()
        pure_content, answer, analysis = _split_answer_analysis(content_no_opts)

        q_type = current_section_type or 'answer'
        if options:
            q_type = 'choice'
        elif '________' in pure_content or '____' in pure_content or '（  ）' in pure_content:
            q_type = 'fill'

        images = _collect_question_images(pure_content, options)
        questions.append({
            'question_number': current_q_num,
            'question_type': q_type,
            'content': pure_content,
            'options': options if options else None,
            'answer': answer,
            'analysis': analysis,
            'difficulty': 3,
            'score': _resolve_question_score(current_q_num, pure_content),
            'images': images if images else None,
        })
        current_q_lines = []

    for line in rich_lines:
        stripped = line.strip()
        if not stripped:
            if current_q_num is not None:
                current_q_lines.append('')
            continue

        # 题干区内若误扫到附录标题，立即停止（双保险）
        if started and _is_appendix_start_line(stripped):
            break

        sec_type, sec_score, sec_per_number = _detect_section(stripped)
        if sec_type is not None:
            _flush_question()
            current_section_type = sec_type
            current_section_score = sec_score
            current_section_per_number = sec_per_number or {}
            started = True
            if sec_per_number:
                logger.info(f"大题类型={sec_type}，题号分值表={sec_per_number}")
            elif sec_score is not None:
                logger.info(f"大题类型={sec_type}，读到每小题分值={sec_score}")
            continue

        q_num = _detect_question_num(stripped)
        if q_num is not None:
            if not started:
                if _is_skip_line(stripped):
                    continue
                after = _strip_question_num(stripped)
                if not after or len(after) < 5:
                    continue
                if '（' not in after and '(' not in after and '?' not in after and '？' not in after:
                    continue

            # 题干区内题号回绕：视为附录开始，停止收题
            if started and last_q_num > 0 and q_num < last_q_num:
                break

            _flush_question()
            current_q_num = q_num
            last_q_num = q_num
            after_num = _strip_question_num(stripped)
            if after_num:
                current_q_lines.append(after_num)
            continue

        if not started:
            continue
        if current_q_num is None:
            continue

        current_q_lines.append(stripped)

    _flush_question()
    return questions


def _parse_questions_from_lines(rich_lines: List[str]) -> List[Dict[str, Any]]:
    """从富文本行列表中解析结构化题目。

    支持两种试卷版式：
    - 题答合一：题干后紧跟【答案】/【解析】（近年常见）
    - 上下分卷：上半仅题干，下半「参考答案」按题号附答案/解析（如杭州2021）
    """
    split_at = _find_split_appendix_index(rich_lines)
    if split_at is not None:
        logger.info(f"检测到上下分卷格式：题干 0..{split_at - 1}，答案附录 {split_at}..")
        questions = _parse_stem_region(rich_lines[:split_at])
        _merge_answer_appendix_lines(questions, rich_lines[split_at:])
    else:
        questions = _parse_stem_region(rich_lines)

    if not questions and any(l.strip() for l in rich_lines):
        full = '\n'.join(l for l in rich_lines if l.strip())
        questions.append({
            'question_number': 1,
            'question_type': 'answer',
            'content': full,
            'options': None,
            'answer': None,
            'analysis': None,
            'difficulty': 3,
            'images': None,
        })

    return questions
