"""Debug image_map for para 12/13"""
from docx import Document
from app.services.word_parser import _extract_images, _extract_para_rich
import os

doc = Document('./uploads/papers/353f79c88c204843b81caf1810895e80.docx')
img_dir = './test_imgs'
os.makedirs(img_dir, exist_ok=True)

image_map = _extract_images(doc, img_dir)
print(f"Total images extracted: {len(image_map)}")
print(f"First 20 rIds: {sorted(list(image_map.keys()))[:20]}")
print(f"First 20 filenames: {list(image_map.values())[:20]}")

print("\n--- Para 12 ---")
print(_extract_para_rich(doc.paragraphs[12], image_map))
print("\n--- Para 13 ---")
print(_extract_para_rich(doc.paragraphs[13], image_map))
print("\n--- Para 14 ---")
print(_extract_para_rich(doc.paragraphs[14], image_map))
